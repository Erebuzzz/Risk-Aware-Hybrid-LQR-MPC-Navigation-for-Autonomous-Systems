from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = "docs/Adaptive_LQR_MPC_Safety_Filter_Idea.docx"


COLORS = {
    "blue": RGBColor(46, 116, 181),
    "dark_blue": RGBColor(31, 77, 120),
    "ink": RGBColor(20, 31, 45),
    "muted": RGBColor(85, 85, 85),
    "light": "F4F6F9",
    "table": "E8EEF5",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")


def set_para_spacing(paragraph, before=0, after=8, line=1.25):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, after=4, line=1.0)
    r = p.add_run("LMS-Adaptive MPC Safety Filter with LQR Nominal Tracking")
    r.font.name = "Calibri"
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = COLORS["ink"]

    p = doc.add_paragraph()
    set_para_spacing(p, after=14, line=1.15)
    r = p.add_run(
        "A publication-oriented research concept for obstacle-aware TurtleBot3 navigation"
    )
    r.font.size = Pt(12)
    r.font.color.rgb = COLORS["muted"]

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [Inches(1.45), Inches(5.05)])
    rows = [
        ("Recommended title", "Risk-Triggered Adaptive MPC Safety Filtering with LQR Nominal Control for Differential-Drive Robots"),
        ("Core claim", "LQR is used as the efficient nominal tracker; adaptive MPC is used only as a safety filter that minimally modifies unsafe LQR commands."),
        ("Best use", "A short conference paper or initial journal submission with simulation plus Gazebo/TurtleBot3 validation."),
    ]
    for i, (label, detail) in enumerate(rows):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = detail
        set_cell_shading(table.cell(i, 0), COLORS["table"])
        for p in table.cell(i, 0).paragraphs:
            p.runs[0].font.bold = True
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                set_para_spacing(p, after=0, line=1.15)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_para_spacing(p, before=14 if level == 1 else 8, after=6, line=1.1)
    r = p.add_run(text)
    r.font.bold = True
    r.font.color.rgb = COLORS["blue"] if level == 1 else COLORS["dark_blue"]
    r.font.size = Pt(16 if level == 1 else 13)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, after=8, line=1.25)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.font.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_para_spacing(p, after=4, line=1.15)
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_para_spacing(p, after=4, line=1.15)
    p.add_run(text)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [Inches(6.5)])
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLORS["light"])
    cell.text = ""
    p = cell.paragraphs[0]
    set_para_spacing(p, after=3, line=1.15)
    r = p.add_run(title)
    r.font.bold = True
    r.font.color.rgb = COLORS["dark_blue"]
    p = cell.add_paragraph()
    set_para_spacing(p, after=0, line=1.15)
    p.add_run(body)
    doc.add_paragraph()


def add_algorithm_table(doc):
    add_heading(doc, "Proposed Control Algorithm", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_width(table, [Inches(0.75), Inches(2.25), Inches(3.5)])
    headers = ["Step", "Module", "Action"]
    for j, h in enumerate(headers):
        table.cell(0, j).text = h
        set_cell_shading(table.cell(0, j), COLORS["table"])
        table.cell(0, j).paragraphs[0].runs[0].font.bold = True
    rows = [
        ("1", "LMS adaptation", "Estimate velocity scale parameters theta_hat = [theta_v, theta_w] from measured state error."),
        ("2", "Adaptive LQR", "Compute the nominal tracking command u_lqr using the current adapted model."),
        ("3", "Risk prediction", "Roll out the LQR command over a short horizon and evaluate distance, time-to-collision, and predicted violation risk."),
        ("4", "Safety filter", "If predicted safe, apply u_lqr. If unsafe, solve adaptive MPC that minimizes deviation from u_lqr while satisfying obstacle constraints."),
        ("5", "Backup safety", "If adaptive MPC is infeasible or too slow, apply a CBF-style brake/turn-away controller."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cells[j].text = value
            for p in cells[j].paragraphs:
                set_para_spacing(p, after=0, line=1.12)
    doc.add_paragraph()


def add_experiment_table(doc):
    add_heading(doc, "Evaluation Plan", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_width(table, [Inches(1.55), Inches(2.4), Inches(2.55)])
    headers = ["Controller", "Purpose", "Expected result"]
    for j, h in enumerate(headers):
        table.cell(0, j).text = h
        set_cell_shading(table.cell(0, j), COLORS["table"])
        table.cell(0, j).paragraphs[0].runs[0].font.bold = True
    rows = [
        ("LQR only", "Nominal tracking baseline", "Low compute cost but obstacle collisions under clutter."),
        ("MPC only", "Constraint-aware baseline", "Better clearance, higher solve time, possible tracking lag."),
        ("Adaptive MPC only", "Model-uncertainty baseline", "Robust tracking but expensive if run continuously."),
        ("Old blended hybrid", "Ablation against current idea", "Smooth commands but no formal safety preservation under blending."),
        ("Proposed method", "Final architecture", "Low intervention rate, fewer collisions, better real-time behavior than continuous adaptive MPC."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cells[j].text = value
            for p in cells[j].paragraphs:
                set_para_spacing(p, after=0, line=1.12)
    doc.add_paragraph()


def add_references(doc):
    add_heading(doc, "Starter References", 1)
    refs = [
        ("Predictive safety filtering", "Wabersich and Zeilinger, predictive safety filter concept.", "https://www.research-collection.ethz.ch/handle/20.500.11850/478766?show=full"),
        ("Control barrier functions", "Ames et al., CBF-QP framework for safety-critical control.", "https://arxiv.org/abs/1609.06408"),
        ("Adaptive MPC for uncertain systems", "Koohler, certainty-equivalent adaptive MPC for uncertain nonlinear systems.", "https://arxiv.org/abs/2603.17843"),
        ("MPC stability and terminal LQR", "Rawlings, Mayne, and Diehl, Model Predictive Control: Theory, Computation, and Design.", "https://sites.engineering.ucsb.edu/~jbraw/mpc/"),
        ("Mobile robot adaptive MPC example", "Adaptive MPC with localization fluctuation for mobile robot trajectory tracking.", "https://pmc.ncbi.nlm.nih.gov/articles/PMC10006979/"),
    ]
    for label, desc, url in refs:
        p = doc.add_paragraph()
        set_para_spacing(p, after=5, line=1.15)
        r = p.add_run(label + ": ")
        r.font.bold = True
        p.add_run(desc + " ")
        add_hyperlink(p, url, url)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLORS["ink"]
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    add_title(doc)

    add_heading(doc, "One-Sentence Research Idea", 1)
    add_callout(
        doc,
        "Main idea",
        "Use LQR as the cheap nominal tracker, then place an LMS-adaptive nonlinear MPC safety filter around it. The adaptive MPC is not a second controller to blend randomly; it is a least-invasive safety layer that changes the LQR command only when predicted future safety constraints are at risk.",
    )

    add_heading(doc, "Why This Is Better Than the Current Hybrid Logic", 1)
    add_body(
        doc,
        "The current LQR-MPC blending idea is difficult to defend because averaging two commands does not guarantee obstacle avoidance. Even if the MPC command is safe, a convex blend with an unsafe nominal command may leave the safety set. This is likely why the current Gazebo behavior feels unreliable.",
    )
    add_body(
        doc,
        "The proposed method changes the framing. LQR is not switched against MPC. LQR is the nominal controller, and adaptive MPC is a predictive safety filter. Reviewers can understand this structure because each part has a clear role: tracking, adaptation, prediction, constraint handling, and fallback safety.",
    )

    add_heading(doc, "Controller Architecture", 1)
    add_body(doc, "The proposed architecture has five layers:")
    add_number(doc, "Reference trajectory or waypoint generator produces x_ref and u_ref.")
    add_number(doc, "LMS estimator updates theta_hat for velocity and yaw-rate scale mismatch.")
    add_number(doc, "Adaptive LQR computes u_lqr using the current model estimate.")
    add_number(doc, "Risk predictor checks whether rolling out u_lqr will violate obstacle clearance.")
    add_number(doc, "Adaptive MPC safety filter minimally modifies u_lqr only when the predicted LQR rollout is unsafe.")

    add_algorithm_table(doc)

    add_heading(doc, "Mathematical Formulation", 1)
    add_body(
        doc,
        "Robot model with online velocity-scale uncertainty:",
    )
    add_body(
        doc,
        "x_dot = theta_v v cos(theta),    y_dot = theta_v v sin(theta),    theta_dot = theta_w omega",
    )
    add_body(
        doc,
        "LMS adaptation updates theta_hat from one-step prediction error:",
    )
    add_body(
        doc,
        "theta_hat(k+1) = clip(theta_hat(k) + Gamma Phi(k)^T e(k), theta_min, theta_max)",
    )
    add_body(
        doc,
        "Safety-filter MPC objective:",
    )
    add_body(
        doc,
        "minimize sum ||x_i - x_ref_i||_Q^2 + ||u_i - u_lqr_i||_R^2 + ||Delta u_i||_S^2",
    )
    add_body(
        doc,
        "subject to adapted nonlinear dynamics, input limits, and ||p_i - p_obs_j|| >= r_obs_j + d_safe for all predicted states and obstacles.",
    )

    add_heading(doc, "What Is Potentially New", 1)
    add_bullet(doc, "Risk-triggered activation: adaptive MPC is solved only when the predicted LQR rollout is unsafe.")
    add_bullet(doc, "Least-invasive correction: MPC minimizes deviation from the LQR command, rather than replacing LQR completely.")
    add_bullet(doc, "Online LMS adaptation: the safety filter predicts using an estimated actuation model, improving behavior under wheel slip, battery sag, or model mismatch.")
    add_bullet(doc, "Safety fallback: a CBF-style emergency controller handles infeasible or slow MPC solves.")
    add_bullet(doc, "Practical validation: compare simulation and Gazebo/TurtleBot3 behavior under obstacle clutter and model mismatch.")

    add_experiment_table(doc)

    add_heading(doc, "Metrics to Report", 1)
    metrics = [
        "Collision count and collision rate",
        "Minimum obstacle clearance",
        "Tracking RMSE and final tracking error",
        "Average and 95th percentile solve time",
        "MPC intervention rate, meaning percent of time MPC had to override LQR",
        "Control smoothness: RMS jerk or command variation",
        "Parameter estimation error for theta_v and theta_w",
    ]
    for item in metrics:
        add_bullet(doc, item)

    add_heading(doc, "Paper Structure", 1)
    for item in [
        "Introduction: why naive hybrid switching/blending is unsafe for obstacle-aware navigation.",
        "Related work: LQR tracking, MPC safety filters, adaptive MPC, CBF backups.",
        "Method: LMS-adaptive LQR nominal controller plus adaptive MPC safety filter.",
        "Experiments: baseline comparisons across static obstacles, dense clutter, and model mismatch.",
        "Results: safety, tracking, compute cost, intervention rate, and parameter adaptation.",
        "Conclusion: the architecture keeps LQR efficiency while adding predictive safety under uncertainty.",
    ]:
        add_number(doc, item)

    add_heading(doc, "Implementation Roadmap for This Project", 1)
    add_body(
        doc,
        "The shortest implementation path is to reuse the existing AdaptiveMPCController and LQRController, but change the outer logic. The new node should compute u_lqr first, roll it out for risk, and call adaptive MPC only when risk is high. The adaptive MPC cost should include ||u - u_lqr|| so it behaves as a safety filter.",
    )
    add_bullet(doc, "Step 1: add a rollout-based risk checker for u_lqr.")
    add_bullet(doc, "Step 2: add an AdaptiveMPCSafetyFilter wrapper around the existing adaptive MPC.")
    add_bullet(doc, "Step 3: add CBF or brake-turn fallback for infeasible solves.")
    add_bullet(doc, "Step 4: run five-controller comparisons and save plots/tables.")
    add_bullet(doc, "Step 5: write the paper around safety filtering, not command blending.")

    add_references(doc)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(footer, after=0, line=1.0)
    run = footer.add_run("Adaptive LQR-MPC Research Concept Brief")
    run.font.size = Pt(9)
    run.font.color.rgb = COLORS["muted"]

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
